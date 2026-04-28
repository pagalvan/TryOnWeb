from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

doc = Document(r'C:\Users\Pablo\Documents\completar tablas con la info.docx')

# ===== DATOS DE LOS 12 CASOS =====
casos = [
    {
        "id": "CA001",
        "requisito": "RF01: Inicio de sesión",
        "escenario": "Ingreso exitoso al sistema como administrador",
        "datos": 'email= "***"\\npass= "***"',
        "pasos": "1. Abrir https://try-on-web.vercel.app/login\n2. Ingresar email en campo \"Email\"\n3. Ingresar contraseña en campo \"Contraseña\"\n4. Click en \"Iniciar Sesión\"",
        "esperado": "Redirección exitosa al dashboard (/dashboard). Se muestra el panel de administración.",
        "obtenido": "Acceso correcto. Redirección al dashboard en 5.117s."
    },
    {
        "id": "CA002",
        "requisito": "RF01: Inicio de sesión",
        "escenario": "Intento de acceso con credenciales incorrectas",
        "datos": 'email= "usuario@fake.com"\npass= "wrongpass"',
        "pasos": "1. Abrir https://try-on-web.vercel.app/login\n2. Ingresar email incorrecto\n3. Ingresar contraseña incorrecta\n4. Click en \"Iniciar Sesión\"",
        "esperado": "Se muestra mensaje de error en rojo. No se redirige al dashboard.",
        "obtenido": "Mensaje de error mostrado correctamente. No hubo redirección. Prueba exitosa en 3.261s."
    },
    {
        "id": "CA003",
        "requisito": "RF02: Registro de usuario",
        "escenario": "Creación exitosa de cuenta nueva de cliente",
        "datos": 'nombre= "Prueba Aceptacion"\ntelefono= "3001234567"\nemail= "prueba.aceptacion2026@gmail.com"\npass= "Test1234"\nconfirmPass= "Test1234"',
        "pasos": "1. Abrir https://try-on-web.vercel.app/registro\n2. Llenar campo \"Nombre completo\"\n3. Llenar campo \"Teléfono\"\n4. Llenar campo \"Correo electrónico\"\n5. Llenar campo \"Contraseña\"\n6. Llenar campo \"Confirmar contraseña\"\n7. Click en \"Crear cuenta\"",
        "esperado": "Toast de confirmación \"Registro completado\". Redirección a /login.",
        "obtenido": "Registro completado exitosamente. Redirección a login en 4.931s."
    },
    {
        "id": "CA004",
        "requisito": "RF03: Recuperación de contraseña",
        "escenario": "Solicitud de restablecimiento de contraseña",
        "datos": 'email= "***"\\npass= "wrongpass"',
        "pasos": "1. Abrir https://try-on-web.vercel.app/login\n2. Ingresar email válido\n3. Ingresar contraseña incorrecta\n4. Click en \"Iniciar Sesión\" (falla)\n5. Click en \"¿Olvidaste tu contraseña?\"",
        "esperado": "Mensaje informativo sobre el envío de correo de restablecimiento.",
        "obtenido": "Se muestra mensaje de recuperación correctamente. Prueba exitosa en 4.10s."
    },
    {
        "id": "CA005",
        "requisito": "RF04: Visualización de productos",
        "escenario": "Filtrar productos por categoría en el catálogo",
        "datos": "Categoría seleccionada = primera categoría disponible en filtros laterales",
        "pasos": "1. Abrir https://try-on-web.vercel.app/productos\n2. Esperar carga de productos\n3. En panel lateral, seleccionar una categoría\n4. Verificar productos filtrados",
        "esperado": "Se muestran solo los productos de la categoría seleccionada.",
        "obtenido": "Filtro de categoría aplicado correctamente. Productos filtrados en 2.988s."
    },
    {
        "id": "CA006",
        "requisito": "RF05: Detalle de producto",
        "escenario": "Acceder a la ficha completa de un producto",
        "datos": "Click en el primer producto visible del catálogo",
        "pasos": "1. Abrir https://try-on-web.vercel.app/productos\n2. Click en el primer producto del grid\n3. Esperar carga de la página de detalle",
        "esperado": "Se muestra página con imagen, nombre, precio, descripción y botones de acción (AR, Probador IA).",
        "obtenido": "Página de detalle cargada correctamente con todos los elementos. Prueba exitosa en 2.410s."
    },
    {
        "id": "CA007",
        "requisito": "RF06: Búsqueda de productos",
        "escenario": "Buscar un producto por nombre",
        "datos": "query = nombre de un producto existente en el catálogo",
        "pasos": "1. Abrir https://try-on-web.vercel.app/productos\n2. Usar la barra de búsqueda del navbar\n3. Escribir parte del nombre del producto\n4. Verificar resultados filtrados",
        "esperado": "El catálogo filtra y muestra solo productos cuyo nombre coincide con la búsqueda.",
        "obtenido": "Búsqueda funcional. Resultados filtrados correctamente en 3.138s."
    },
    {
        "id": "CA008",
        "requisito": "RF07: Probador Virtual IA",
        "escenario": "Abrir el probador virtual con IA (Gemini) desde un producto",
        "datos": "Producto con imagen disponible, usuario autenticado",
        "pasos": "1. Login con usuario válido\n2. Ir a /productos\n3. Seleccionar un producto\n4. Click en botón \"Probador IA\"\n5. Verificar apertura del diálogo",
        "esperado": "Se abre el diálogo \"Probador Virtual con Gemini IA\" con zonas de \"Tu Foto\" y \"Prenda\".",
        "obtenido": "Diálogo de Probador IA abierto correctamente. Interfaz funcional en 3.453s."
    },
    {
        "id": "CA009",
        "requisito": "RF08: CRUD Productos",
        "escenario": "Crear un nuevo producto desde el panel de inventario",
        "datos": 'nombre= "Producto Test CA009"\nsku= generado automáticamente\ncategoría= primera disponible\nprecio= "50000"\ndescripcion= "Producto de prueba"',
        "pasos": "1. Login como admin\n2. Navegar a /inventario\n3. Click en \"Nuevo Producto\"\n4. Llenar nombre, SKU (auto), categoría, precio, descripción\n5. Click en \"Crear\"",
        "esperado": "Toast \"Producto creado\". El producto aparece en la tabla de inventario.",
        "obtenido": "Producto creado exitosamente. Visible en tabla de inventario. Prueba exitosa en 33.905s."
    },
    {
        "id": "CA010",
        "requisito": "RF09: Gestión de Stock",
        "escenario": "Modificar el stock de un producto existente",
        "datos": 'producto= primer producto en tabla\nnueva cantidad= "25"',
        "pasos": "1. Login como admin\n2. Navegar a /inventario\n3. Click en menú (⋮) del primer producto\n4. Click en \"Gestionar stock\"\n5. Cambiar cantidad a 25\n6. Click en \"Guardar\"",
        "esperado": "Toast \"Stock actualizado\". La cantidad del producto se actualiza a 25.",
        "obtenido": "Stock actualizado correctamente a 25 unidades. Prueba exitosa en 30.432s."
    },
    {
        "id": "CA011",
        "requisito": "RF10: CRUD Categorías",
        "escenario": "Crear una nueva categoría de productos",
        "datos": 'nombre= "Categoria Test CA011"\ndescripcion= "Categoria de prueba"\nicono= primer icono disponible',
        "pasos": "1. Login como admin\n2. Navegar a /categorias\n3. Click en \"Nueva Categoría\"\n4. Llenar nombre, descripción, seleccionar icono\n5. Click en \"Crear\"",
        "esperado": "Toast \"Categoría creada\". La categoría aparece en el grid con estado \"Activa\".",
        "obtenido": "Categoría creada exitosamente. Visible en grid de categorías. Prueba exitosa."
    },
    {
        "id": "CA012",
        "requisito": "RF11: Dashboard analítico",
        "escenario": "Acceder al panel de estadísticas del administrador",
        "datos": "Usuario admin autenticado",
        "pasos": "1. Login como admin\n2. Verificar redirección automática a /dashboard\n3. Verificar secciones del panel",
        "esperado": "Se muestra el dashboard con secciones de resumen, navegación a sub-secciones (Alertas, Analytics, Catálogo, Demanda, Inventario, Operaciones, Probador).",
        "obtenido": "Dashboard cargado correctamente con todas las secciones. Prueba exitosa en 8.180s."
    },
]

# ===== LIMPIAR DOCUMENTO Y RECONSTRUIR =====
# Eliminar tablas existentes (excepto mantener estructura del doc)
while len(doc.tables) > 0:
    table = doc.tables[0]
    table._element.getparent().remove(table._element)

# Eliminar párrafos existentes
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)

# ===== RECONSTRUIR DOCUMENTO =====

# Título principal
h1 = doc.add_heading('PRUEBAS DE ACEPTACIÓN', level=1)

# Subtítulo
h2 = doc.add_heading('Diseño de los casos de pruebas', level=2)

# Crear una tabla por cada caso de prueba
for caso in casos:
    table = doc.add_table(rows=3, cols=3)
    # table border style applied by default
    
    # Fila 1: ID | Requisito | Escenario
    table.cell(0, 0).text = f"ID Caso: {caso['id']}"
    table.cell(0, 1).text = f"Requisito Asociado: {caso['requisito']}"
    table.cell(0, 2).text = f"Escenario de Negocio:\n{caso['escenario']}"
    
    # Fila 2: Datos | Pasos (merged)
    table.cell(1, 0).text = f"Datos de entrada:\n{caso['datos']}"
    # Merge cells 1,1 and 1,2 for pasos
    cell_pasos = table.cell(1, 1)
    cell_pasos.merge(table.cell(1, 2))
    cell_pasos.text = f"Pasos de Ejecución:\n{caso['pasos']}"
    
    # Fila 3: Resultado Esperado | Resultado Obtenido (merged)
    table.cell(2, 0).text = f"Resultado Esperado:\n{caso['esperado']}"
    cell_obtenido = table.cell(2, 1)
    cell_obtenido.merge(table.cell(2, 2))
    cell_obtenido.text = f"Resultado Obtenido: {caso['obtenido']}"
    
    # Espacio entre tablas
    doc.add_paragraph('')

# ===== SECCIÓN: Ejecución y evaluación =====
doc.add_heading('Ejecución y evaluación de las pruebas', level=2)
doc.add_paragraph('Las pruebas fueron ejecutadas con Katalon Studio 11.1.2 (Community Edition) sobre el navegador Chrome 147.0.7727.102 en Windows 10 64bit.')
doc.add_paragraph('URL de pruebas: https://try-on-web.vercel.app/')
doc.add_paragraph('Fecha de ejecución: 2026-04-26')
doc.add_paragraph('Duración total: 2m 15.739s')
doc.add_paragraph('')

# Tabla de evaluación
eval_table = doc.add_table(rows=len(casos) + 1, cols=3)
# eval_table.style = 'Table Grid'

# Header
eval_table.cell(0, 0).text = "Caso de prueba"
eval_table.cell(0, 1).text = "Resultado"
eval_table.cell(0, 2).text = "Observaciones"

# Hacer header bold
for cell in eval_table.rows[0].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Filas de datos
for i, caso in enumerate(casos):
    eval_table.cell(i + 1, 0).text = f"{caso['id']}: {caso['escenario']}"
    eval_table.cell(i + 1, 1).text = "✅ Passed"
    eval_table.cell(i + 1, 2).text = caso['obtenido']

doc.add_paragraph('')

# ===== Observaciones Técnicas =====
doc.add_heading('Observaciones Técnicas/Funcionales', level=3)
doc.add_paragraph('• Los 12 casos de prueba automatizados con Katalon Studio fueron ejecutados satisfactoriamente con resultado PASSED.')
doc.add_paragraph('• Los casos de prueba del panel de administración (CA009, CA010, CA011) requirieron scripts personalizados en Groovy debido a que el Web Recorder de Katalon no captura correctamente los componentes dinámicos de Radix UI (diálogos modales, selects, dropdowns).')
doc.add_paragraph('• Se detectó y corrigió un bug en la funcionalidad de recuperación de contraseña (CA004): el sistema no mostraba feedback al usuario cuando Supabase respondía con error de rate limit. Se implementó manejo de errores con mensajes amigables.')
doc.add_paragraph('• La integración del Probador Virtual con IA (Gemini) funciona correctamente desde la interfaz web. Las pruebas de AR con Snap Camera Kit no se automatizaron por requerir hardware de cámara.')

# ===== Concepto Final =====
doc.add_heading('Concepto Final', level=3)
doc.add_paragraph('El aplicativo TryOnWeb cumple satisfactoriamente con los requisitos funcionales evaluados. Los 12 casos de prueba de aceptación fueron ejecutados con resultado exitoso (12/12 PASSED). El sistema es estable, funcional y está listo para uso en producción.')

# ===== CONCLUSIONES =====
doc.add_heading('CONCLUSIONES', level=1)

doc.add_heading('Pruebas de Autenticación (CA001-CA004)', level=3)
doc.add_paragraph('El módulo de autenticación del sistema TryOnWeb demostró un comportamiento robusto y confiable. El inicio de sesión exitoso redirige correctamente al dashboard, el manejo de credenciales incorrectas muestra mensajes de error apropiados, el registro de nuevos usuarios funciona con validación de campos, y la recuperación de contraseña opera correctamente con manejo de rate limit de Supabase. Se concluye que el flujo de autenticación cumple con los requisitos funcionales establecidos.')

doc.add_heading('Pruebas de Catálogo (CA005-CA007)', level=3)
doc.add_paragraph('El catálogo de productos presenta una experiencia de usuario fluida y funcional. Los filtros por categoría operan correctamente actualizando la vista en tiempo real, la página de detalle de producto muestra toda la información relevante incluyendo botones de acción para AR y Probador IA, y la búsqueda por nombre filtra los resultados de forma precisa. Se concluye que la experiencia del cliente en el catálogo es satisfactoria.')

doc.add_heading('Pruebas del Probador Virtual (CA008)', level=3)
doc.add_paragraph('La funcionalidad de Probador Virtual con IA (Gemini) se integra correctamente desde la ficha de producto. El diálogo se abre con las zonas de "Tu Foto" y "Prenda" operativas. Se concluye que la funcionalidad de probador virtual cumple con el requisito funcional, aunque las pruebas de AR con cámara real requieren validación manual adicional.')

doc.add_heading('Pruebas de Inventario (CA009-CA010)', level=3)
doc.add_paragraph('El módulo de gestión de inventario permite crear productos con todos sus campos (nombre, SKU auto-generado, categoría, precio, descripción) y gestionar stock por bodega de forma eficiente. La actualización de cantidades se refleja inmediatamente en la tabla. Se concluye que el CRUD de productos y la gestión de stock cumplen con los requisitos administrativos del sistema.')

doc.add_heading('Pruebas de Categorías (CA011)', level=3)
doc.add_paragraph('La creación de categorías funciona correctamente con asignación de nombre, descripción e icono. Las nuevas categorías aparecen inmediatamente en el grid con estado "Activa". Se concluye que la gestión de categorías opera según lo esperado.')

doc.add_heading('Pruebas de Dashboard (CA012)', level=3)
doc.add_paragraph('El dashboard administrativo carga correctamente con todas sus secciones de navegación (Alertas, Analytics, Catálogo, Demanda, Inventario, Operaciones, Probador). Se concluye que el panel analítico proporciona la visibilidad necesaria para la gestión del sistema.')

# ===== SECCIÓN C: Ejecución con herramienta =====
doc.add_heading('Ejecución y evaluación de las pruebas', level=1)
doc.add_paragraph('Las pruebas fueron ejecutadas con Katalon Studio 11.1.2 (Community Edition).')
doc.add_paragraph('Herramienta: Katalon Studio')
doc.add_paragraph('Navegador: Chrome 147.0.7727.102')
doc.add_paragraph('Plataforma: Windows 10 64bit')
doc.add_paragraph('Total de casos: 12')
doc.add_paragraph('Resultado: 12 Passed, 0 Failed, 0 Error')
doc.add_paragraph('Duración total: 2m 15.739s')
doc.add_paragraph('')
doc.add_paragraph('[Insertar aquí las capturas de pantalla del reporte de Katalon Studio y de la ejecución paso a paso de cada caso de prueba]')

# Guardar
output_path = r'C:\Users\Pablo\Documents\Pruebas_Aceptacion_TryOnWeb_Completo.docx'
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
